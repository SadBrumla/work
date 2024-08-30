import os
from docx import Document
from docx.shared import Mm
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

def save_to_doc(doc, path, name):
        doc.add_picture(path + '/' + name, width=Mm(155), height=Mm(100))
        pass

def name_setting(file_name, ventily, x_val):
    x_val_dict =    {"GH":"u/cᵢₛ",
                     "GL":"u/cᵢₛ",
                     "FC":"Ma",
                     "EZ":"Ma"}
    x_description = x_val_dict[x_val]

    ID_list = ["3-59","104-139","bez_premosteni_200-302","s_premostenim_200-302","312-359","475-510","473,474,526,527","540-588","1-89"]
    ID = [s for s in ID_list if s in file_name][0]

    if ID == "bez_premosteni_200-302":
        ID = "200-302 (bez přemostění)"
    elif ID == "s_premostenim_200-302":
        ID = "200-302 (s přemostěním)"

    string = f"Měření {ID}, {ventily}, {x_description}"
    return string




def name_figure(i,x_val, y_val, ventily):
    x_val_dict =    {"GH":"u/cᵢₛ",
                     "GL":"u/cᵢₛ",
                     "FC":"Ma",
                     "EZ":"Ma"}
    x_description = x_val_dict[x_val]

    y_val_list = ["_p_0","_t_0","G_rl","Mk_","N_nl","ro_s","Ma_sRL","Re_sRL","eta_"]
    y_index =next((j for j, s in enumerate(y_val_list) if s in y_val), -1)
    y_descriptions = ["Tlaky v jednotlivých rovinách",
                      "Teploty v jednotlivých rovinách",
                      "Průtok RL obou stupňů",
                      "Kroutící moment z brzd",
                      "Výkony na lopatkách obou stupňů",
                      "Reakce obou stupňů",
                      "Machova čísla RL",
                      "Reynoldsova čísla RL",
                      "Účinnosti z brzdy"]

    y_description = y_descriptions[y_index]

    string = f"Obr. {i}: {y_description} - {ventily} - {x_description}"
    i += 1 
    return string

ventily_list = ["2V","3V","4V"]
doc = Document()
dir =  "S:/Rozvoj/01_Experimental_Research/Ranš/T1MW_grafy/Vygenerovane_grafy_T1MW"
setting_list = os.listdir(dir)
order_of_graphs = ["_p_0","_t_0","G_rl","Mk_","N_nl","ro_s","Ma_sRL","Re_sRL","eta_"]
i = 1

# loop thrugh every setting 
for setting in setting_list:

    try:
        ventily = [s for s in ventily_list if s in setting][0]
    except:
        ventily = ""

    

    setting_dir = dir + "/" + setting
    x_vals = os.listdir(setting_dir)

    if x_vals[0] == "EZ":
        x_vals = sorted(x_vals,reverse=True)

    path1 = setting_dir + "/" + x_vals[0]
    path2 = setting_dir + "/" + x_vals[1]

    files1 = os.listdir(path1)
    files2 = os.listdir(path2)


    # loop thrugh all files in x_val directory (all pictures in EZ)
    for graf_start in order_of_graphs:

        graf1_list = [s for s in files1 if graf_start in s]
        graf2_list = [s for s in files2 if graf_start in s]

        setting_name = name_setting(setting,ventily,x_vals[0])

        if len(graf1_list) == 1 and len(graf2_list) == 1 :
            p = doc.add_paragraph(setting_name)
            save_to_doc(doc,path1,name = graf1_list[0])
            save_to_doc(doc,path2,name = graf2_list[0])

        elif len(graf1_list) == 2 and len(graf2_list) == 2 :
            p = doc.add_paragraph(setting_name)
            save_to_doc(doc,path1,name = graf1_list[0])
            save_to_doc(doc,path2,name = graf2_list[1])


        elif len(graf1_list) == 3 and len(graf2_list) == 3 :
            p = doc.add_paragraph(setting_name)
            save_to_doc(doc,path1,name = graf1_list[0])
            save_to_doc(doc,path2,name = graf2_list[1])
            d = doc.add_paragraph(name_figure(i,x_vals[0],graf1_list[0],ventily))
            i += 1
            doc.add_page_break()

            p = doc.add_paragraph(setting_name)
            save_to_doc(doc,path1,name = graf1_list[2])
            save_to_doc(doc,path2,name = graf2_list[2])

        d = doc.add_paragraph(name_figure(i,x_vals[0],graf1_list[0],ventily))
        i += 1
        if graf_start == order_of_graphs[-1]:
            doc.add_section()
        else:
            doc.add_page_break()


j = 0
c = 0

colors = [
    RGBColor(204, 0, 0),      # Dark Red
    RGBColor(0, 102, 0),      # Dark Green
    RGBColor(0, 0, 102),      # Dark Blue
    RGBColor(204, 102, 0),    # Dark Orange
    RGBColor(102, 0, 102),    # Dark Purple    
    RGBColor(0, 102, 102),    # Dark Teal
    RGBColor(204, 102, 153),  # Dark Pink
    RGBColor(51, 51, 153),    # Dark Royal Blue
    RGBColor(153, 153, 0)     # Dark Gold
]


text_header_old = doc.paragraphs[0].text
for para in doc.paragraphs:
    if para.text == "":
        continue

    para_format = para.paragraph_format
    if j%2 == 0:
        if text_header_old != para.text:
            c += 1
        para_format.space_before = Pt(6) 
        para_format.space_after = Pt(6)    
        para_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for run in para.runs:

            run.font.name = 'Calibri'
            run.font.size = Pt(10) 
            run.font.bold = True
            run.font.color.rgb = colors[c]

        text_header_old = para.text

    elif j%2 == 1:
        para_format.space_before = Pt(6) 
        para_format.space_after = Pt(6)    
        para_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10) 
            run.font.bold = False

    j += 1



doc.save('S:\Rozvoj\\01_Experimental_Research\Ranš\T1MW_grafy\grafy.docx')



